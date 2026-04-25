"""CV Tailor — desktop app that tailors CVs to job descriptions via Claude."""
import csv
import json
import os
import re
import subprocess
import sys
import threading
import time
import unicodedata
from datetime import datetime
from pathlib import Path

import keyring
import pdfplumber
from anthropic import (Anthropic, APIConnectionError, AuthenticationError,
                       RateLimitError)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QColor
from PyQt5.QtWidgets import (QApplication, QFileDialog, QFormLayout,
                             QGroupBox, QHBoxLayout, QHeaderView, QLabel,
                             QLineEdit, QListWidget, QListWidgetItem,
                             QMainWindow, QMessageBox, QProgressBar,
                             QPushButton, QTableWidget, QTableWidgetItem,
                             QTabWidget, QTextEdit, QVBoxLayout, QWidget)

APP_VERSION = "1.1.0"
MODEL_NAME = "claude-sonnet-4-5"
MAX_TOKENS_UNIFIED = 3000
DEBUG_CV = True
MAX_TOKENS_PROFILE = 8000
CONFIG_FILE = "config.json"
PROFILE_FILE = "master_profile.json"
DEFAULT_OUTPUT = "output"
KEYRING_SERVICE = "cv-tailor"
KEYRING_USER = "anthropic-api-key"
COMPANY_MAX = 25
TITLE_MAX = 35
BULK_DELAY_SEC = 1.0
RATE_LIMIT_RETRY_SEC = 5

PAGE_W_CM = 21.0
PAGE_H_CM = 29.7
MARGIN_CM = 2.0
RIGHT_TAB_CM = PAGE_W_CM - 2 * MARGIN_CM  # 17cm — dynamic right-edge tab

BG = "#0d1117"
FG = "#e6edf3"
PANEL = "#161b22"
BORDER = "#30363d"
ACCENT = "#388bfd"
GREEN = "#238636"
YELLOW = "#9e6a03"
RED = "#da3633"

PROFILE_SCHEMA = {
    "name": "", "contact": {"email": "", "phone": "", "linkedin": "",
                            "github": "", "website": ""},
    "summary": "", "experience": [], "education": [],
    "skills": {"languages": [], "frontend": [], "backend": [],
               "databases": [], "cloud": [], "ai_integrations": [],
               "third_party_apis": [], "erp": [], "other": []},
    "projects": [], "certifications": [], "volunteering": [],
    "achievements": [], "interests": []
}

ALWAYS_SHOW_SKILLS = [
    "languages", "frontend", "backend", "databases",
    "cloud", "ai_integrations", "third_party_apis",
]

CONDITIONAL_SKILLS = [
    "erp", "desktop_gui", "productivity_tools", "design_collaboration",
    "analytics_tools", "dev_tools", "soft_skills", "languages_spoken",
]

SKILL_LABELS = {
    "languages": "Languages", "frontend": "Frontend",
    "backend": "Backend", "databases": "Databases",
    "cloud": "Cloud & DevOps", "ai_integrations": "AI / Integrations",
    "third_party_apis": "Third-Party APIs", "erp": "ERP",
    "desktop_gui": "Desktop / GUI", "productivity_tools": "Productivity",
    "design_collaboration": "Design & Collab",
    "analytics_tools": "Analytics", "dev_tools": "Dev Tools",
    "soft_skills": "Soft Skills", "languages_spoken": "Languages Spoken",
}

# Ordered list of all categories for profile display
SKILL_CATEGORIES = [(k, v) for k, v in SKILL_LABELS.items()]


# ─── Config / IO helpers ─────────────────────────────────────────────────

def _get_api_key_from_keyring():
    """Read the API key from the OS keyring. Returns '' on any failure."""
    try:
        return keyring.get_password(KEYRING_SERVICE, KEYRING_USER) or ""
    except Exception:
        return ""


def _set_api_key_in_keyring(key):
    """Store (or delete if empty) the API key in the OS keyring."""
    try:
        if key:
            keyring.set_password(KEYRING_SERVICE, KEYRING_USER, key)
        else:
            try:
                keyring.delete_password(KEYRING_SERVICE, KEYRING_USER)
            except Exception:
                pass
    except Exception:
        pass


def load_config():
    """Return config dict. API key comes from OS keyring; if a legacy
    plaintext key is found in config.json it is migrated to the keyring
    and scrubbed from disk."""
    cfg = {"api_key": "", "output_folder": DEFAULT_OUTPUT}
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                disk = json.load(f)
            cfg["output_folder"] = disk.get("output_folder") or DEFAULT_OUTPUT
            legacy_key = (disk.get("api_key") or "").strip()
            if legacy_key:
                _set_api_key_in_keyring(legacy_key)
                # Scrub plaintext key from disk
                try:
                    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                        json.dump({"output_folder": cfg["output_folder"]},
                                  f, indent=2)
                except Exception:
                    pass
        except Exception:
            pass
    cfg["api_key"] = _get_api_key_from_keyring()
    return cfg


def save_config(cfg):
    """Persist non-secret config to disk; persist API key to OS keyring."""
    _set_api_key_in_keyring((cfg.get("api_key") or "").strip())
    on_disk = {"output_folder": cfg.get("output_folder") or DEFAULT_OUTPUT}
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(on_disk, f, indent=2)


def load_profile():
    """Return master profile dict, or None if missing/corrupt."""
    if not os.path.exists(PROFILE_FILE):
        return None
    try:
        with open(PROFILE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def save_profile(profile):
    """Persist master profile to JSON."""
    with open(PROFILE_FILE, "w", encoding="utf-8") as f:
        json.dump(profile, f, indent=2)


def extract_text_from_file(path):
    """Extract raw text from PDF, DOCX, or TXT."""
    ext = Path(path).suffix.lower()
    try:
        if ext == ".pdf":
            with pdfplumber.open(path) as pdf:
                return "\n".join((p.extract_text() or "") for p in pdf.pages)
        if ext == ".docx":
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        if ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        return f"[parse error: {e}]"
    return ""


def sanitize_filename_part(text, maxlen):
    """Turn text into safe filename segment, transliterating non-ASCII."""
    if not text:
        return "Unknown"
    text = unicodedata.normalize("NFKD", text)
    text = text.encode("ascii", "ignore").decode("ascii")
    text = text.strip().replace(" ", "_")
    text = re.sub(r"[^A-Za-z0-9_\-]", "", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return text[:maxlen] or "Unknown"


def build_output_path(output_folder, company, title):
    """Return non-colliding output path for today's CV."""
    os.makedirs(output_folder, exist_ok=True)
    date = datetime.now().strftime("%Y-%m-%d")
    base = f"{date}_{sanitize_filename_part(company, COMPANY_MAX)}_" \
           f"{sanitize_filename_part(title, TITLE_MAX)}"
    path = os.path.join(output_folder, base + ".docx")
    i = 2
    while os.path.exists(path):
        path = os.path.join(output_folder, f"{base}_{i}.docx")
        i += 1
    return path


def trim_profile(profile):
    """Drop empty values so we don't pay tokens for blank keys."""
    if not profile:
        return {}
    out = {}
    for k, v in profile.items():
        if v in (None, "", [], {}):
            continue
        if isinstance(v, dict):
            sub = {sk: sv for sk, sv in v.items() if sv not in (None, "", [])}
            if sub:
                out[k] = sub
        elif isinstance(v, list):
            cleaned = [x for x in v if x not in (None, "", [], {})]
            if cleaned:
                out[k] = cleaned
        else:
            out[k] = v
    return out


_FIELD_SHORTS = {
    "achievements": "ach", "responsibilities": "resp",
    "organization": "org", "description": "desc",
    "technologies": "tech", "highlights": "pts",
}
_DROP_TOP_KEYS = {"interests", "certifications"}
_DROP_EXP_KEYS = {"location"}


def slim_profile_for_generation(profile):
    """Drop noise fields Claude doesn't need. Keeps the cached payload small."""
    if not profile:
        return {}
    out = {}
    for k, v in profile.items():
        if k in _DROP_TOP_KEYS or v in (None, "", [], {}):
            continue
        if k == "experience" and isinstance(v, list):
            out[k] = [
                {ek: ev for ek, ev in exp.items()
                 if ek not in _DROP_EXP_KEYS and ev not in (None, "", [], {})}
                for exp in v if isinstance(exp, dict)
            ]
        elif isinstance(v, dict):
            sub = {sk: sv for sk, sv in v.items()
                   if sv not in (None, "", [], {})}
            if sub:
                out[k] = sub
        elif isinstance(v, list):
            cleaned = []
            for item in v:
                if isinstance(item, dict):
                    ci = {ik: iv for ik, iv in item.items()
                          if iv not in (None, "", [], {})}
                    if ci:
                        cleaned.append(ci)
                elif item not in (None, "", [], {}):
                    cleaned.append(item)
            if cleaned:
                out[k] = cleaned
        else:
            out[k] = v
    return out


def compress_profile(profile):
    """Slim, serialize tight, and shorten common keys to save input tokens."""
    slim = slim_profile_for_generation(profile)
    text = json.dumps(slim, separators=(",", ":"))
    for long, short in _FIELD_SHORTS.items():
        text = text.replace(f'"{long}":', f'"{short}":')
    return text


def validate_cv_output(cv_data, master_profile):
    """Strip any skill or achievement not present in the master profile."""
    real_skills = set()
    for cat_skills in master_profile.get("skills", {}).values():
        if isinstance(cat_skills, list):
            for s in cat_skills:
                real_skills.add(s.lower().strip())
    for cat, skills in cv_data.get("skills", {}).items():
        if isinstance(skills, list):
            cv_data["skills"][cat] = [
                s for s in skills if s.lower().strip() in real_skills
            ]
    real_ach = {a.lower().strip()
                for a in master_profile.get("achievements", [])}
    if real_ach:
        cv_data["achievements"] = [
            a for a in cv_data.get("achievements", [])
            if a.lower().strip() in real_ach
        ]
    return cv_data


def strip_hard_gap(text):
    """Split a response into (body_without_hard_gap, hard_gap_string)."""
    m = re.search(r"(?im)^\s*HARD_GAP\s*:\s*(.+?)\s*$", text)
    if not m:
        return text, ""
    return text[:m.start()] + text[m.end():], m.group(1).strip()


def parse_json_response(raw_text):
    """Extract a JSON object from Claude output robustly."""
    text = raw_text.strip()
    # Strip markdown fences
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
    text = re.sub(r"```[a-zA-Z]*\n?", "", text).replace("```", "")

    # Direct parse
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Boundary extraction: first { to last }
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start:end + 1])
        except json.JSONDecodeError:
            pass

    # Log details for diagnosis
    print(f"[ERROR] Failed to parse Claude response")
    print(f"[ERROR] Response length: {len(raw_text)} chars")
    print(f"[ERROR] First 500 chars: {raw_text[:500]}")
    print(f"[ERROR] Last 200 chars: {raw_text[-200:]}")
    raise ValueError(
        f"Could not parse Claude response as JSON. "
        f"Response length: {len(raw_text)} chars."
    )


def open_file_native(path):
    """Open a file with the OS default application (Windows-first)."""
    try:
        os.startfile(os.path.abspath(path))
    except Exception:
        subprocess.Popen(["cmd", "/c", "start", "", os.path.abspath(path)],
                         shell=False)


def safe_merge_profiles(old, new):
    """Guard against Claude dropping data: keep any old list items missing
    from new (by name/title key)."""
    if not old:
        return new
    if not new:
        return old
    out = dict(new)
    for list_key, id_key in (("experience", "title"), ("projects", "name"),
                             ("education", "degree"),
                             ("certifications", "name"),
                             ("volunteering", "role")):
        old_list = old.get(list_key) or []
        new_list = out.get(list_key) or []
        seen = {(it.get(id_key) or "").strip().lower() for it in new_list}
        for it in old_list:
            key = (it.get(id_key) or "").strip().lower()
            if key and key not in seen:
                new_list.append(it)
                seen.add(key)
        out[list_key] = new_list
    # Union skill lists
    old_sk = old.get("skills") or {}
    new_sk = out.get("skills") or {}
    merged_sk = {}
    for k in set(list(old_sk.keys()) + list(new_sk.keys())):
        merged = list(new_sk.get(k) or [])
        lower = {s.lower() for s in merged}
        for s in old_sk.get(k) or []:
            if s.lower() not in lower:
                merged.append(s)
                lower.add(s.lower())
        merged_sk[k] = merged
    if merged_sk:
        out["skills"] = merged_sk
    # Preserve contact fields that got blanked
    old_c = old.get("contact") or {}
    new_c = out.get("contact") or {}
    for k, v in old_c.items():
        if v and not new_c.get(k):
            new_c[k] = v
    if new_c:
        out["contact"] = new_c
    if old.get("name") and not out.get("name"):
        out["name"] = old["name"]
    return out


# ─── Claude call wrapper with parse-retry + rate-limit retry ─────────────

def _log_usage(msg, text):
    """Emit token + cache stats when DEBUG_CV is on."""
    if not DEBUG_CV:
        return
    u = msg.usage
    cc = getattr(u, "cache_creation_input_tokens", 0) or 0
    cr = getattr(u, "cache_read_input_tokens", 0) or 0
    print(f"[DEBUG] in={u.input_tokens} out={u.output_tokens} "
          f"cache_created={cc} cache_read={cr} "
          f"chars={len(text)} preview={text[:200]!r}")


def claude_call(client, system, user, max_tokens, retries=1):
    """Call Anthropic once, retrying on RateLimit once. Returns raw text."""
    last_err = None
    for _ in range(retries + 2):
        try:
            msg = client.messages.create(
                model=MODEL_NAME, max_tokens=max_tokens, system=system,
                messages=[{"role": "user", "content": user}])
            text = msg.content[0].text
            _log_usage(msg, text)
            return text
        except RateLimitError as e:
            last_err = e
            time.sleep(RATE_LIMIT_RETRY_SEC)
            continue
    raise last_err


def claude_call_cached(client, system, cached_user, fresh_user,
                       max_tokens, retries=1):
    """Like claude_call, but marks system + cached_user for prompt caching.

    Cached blocks cost 10% of the normal input rate on a cache-hit and
    ~125% on cache creation. For the CV generation flow, system prompt
    and profile are static across calls, only the JD changes.
    """
    last_err = None
    for _ in range(retries + 2):
        try:
            msg = client.messages.create(
                model=MODEL_NAME,
                max_tokens=max_tokens,
                system=[{
                    "type": "text",
                    "text": system,
                    "cache_control": {"type": "ephemeral"},
                }],
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": cached_user,
                            "cache_control": {"type": "ephemeral"},
                        },
                        {"type": "text", "text": fresh_user},
                    ],
                }],
            )
            text = msg.content[0].text
            _log_usage(msg, text)
            return text
        except RateLimitError as e:
            last_err = e
            time.sleep(RATE_LIMIT_RETRY_SEC)
            continue
    raise last_err


# ─── ProfileManager ──────────────────────────────────────────────────────

class ProfileManager:
    """Build and merge the master profile via Claude."""

    def __init__(self, client):
        """Hold Anthropic client."""
        self.client = client

    def build_new(self, texts):
        """Create a fresh profile from extracted document texts."""
        combined = "\n\n---\n\n".join(texts)
        schema_str = json.dumps(PROFILE_SCHEMA, separators=(",", ":"))
        system = (
            "Extract ONLY the information explicitly present in the CV "
            "document(s) into this exact JSON schema. Return ONLY valid "
            "JSON, no fences, no commentary.\n\n"
            "STRICT EXTRACTION RULES:\n"
            "- Copy text verbatim where possible; minor rewording only.\n"
            "- NEVER invent metrics, numbers, percentages, dates, awards, "
            "projects, skills, languages, or achievements not stated in "
            "the source documents.\n"
            "- If a field is not in the source, leave it empty.\n"
            "- Experience entries use keys: title, company, location, "
            "start_date, end_date, responsibilities (list of bullets).\n"
            "- Project entries use keys: name, description, technologies "
            "(list), link, highlights (list of bullets).\n"
            "- Volunteering entries use keys: role, organization, "
            "start_date, end_date, description.\n"
            "- Education entries use keys: degree, institution, location, "
            "start_date, end_date, details.\n"
            "- Skills must use the category keys in the schema exactly.\n\n"
            "SCHEMA: " + schema_str
        )
        text = claude_call(self.client, system, combined, MAX_TOKENS_PROFILE)
        return parse_json_response(text)

    def merge(self, existing, texts):
        """Merge new document texts into existing profile."""
        combined = "\n\n---\n\n".join(texts)
        compact = json.dumps(existing, separators=(",", ":"))
        system = (
            "Merge new CV information into the existing profile. Add new "
            "roles, projects, skills found in the source documents. "
            "Deduplicate. NEVER remove existing data. NEVER invent "
            "content not present in either the existing profile or the "
            "new documents (no fabricated metrics, awards, projects, "
            "skills, or achievements). Use the same schema keys as the "
            "existing profile. Return the complete updated profile as "
            "JSON only, no fences."
        )
        user = f"EXISTING:{compact}\n\nNEW:\n{combined}"
        text = claude_call(self.client, system, user, MAX_TOKENS_PROFILE)
        merged = parse_json_response(text)
        return safe_merge_profiles(existing, merged)


# ─── DocxBuilder ─────────────────────────────────────────────────────────

class DocxBuilder:
    """Produce a formatted .docx CV from tailored JSON data."""

    @staticmethod
    def _set_bottom_border(paragraph):
        """Add a single-line bottom border to a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _clear_table_borders(table):
        """Force all borders off (top/left/bottom/right/insideH/insideV)."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tblPr.append(borders)

    @staticmethod
    def _set_cell_width(cell, cm):
        """Set a fixed cell width in OXML (python-docx cell.width is lossy)."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(int(cm * 567)))  # 567 twips per cm
        tcW.set(qn("w:type"), "dxa")

    @staticmethod
    def _add_run(paragraph, text, *, bold=False, italic=False,
                 size=10.5, color=None, font="Calibri"):
        """Append a styled run to a paragraph."""
        run = paragraph.add_run(text)
        run.font.name = font
        # Needed for East-Asian fallback to also use Calibri in Word
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rFonts.set(qn("w:cs"), font)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        return run

    @staticmethod
    def _set_defaults(doc):
        """Set document default font to Calibri 10.5 so styles inherit."""
        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        rPr = normal.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        rFonts.set(qn("w:cs"), "Calibri")

    @staticmethod
    def _set_page(doc):
        """Set A4, 2cm margins."""
        for section in doc.sections:
            section.page_width = Cm(PAGE_W_CM)
            section.page_height = Cm(PAGE_H_CM)
            section.left_margin = Cm(MARGIN_CM)
            section.right_margin = Cm(MARGIN_CM)
            section.top_margin = Cm(MARGIN_CM)
            section.bottom_margin = Cm(MARGIN_CM)

    @classmethod
    def _section_header(cls, doc, text):
        """Add an ALL CAPS section header with bottom border."""
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        cls._add_run(p, text.upper(), bold=True, size=11, color="000000")
        cls._set_bottom_border(p)
        return p

    @classmethod
    def _title_with_date(cls, doc, title, date_str, italic_title=False):
        """Bold title left, date right, right-tab at (page − margins)."""
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(RIGHT_TAB_CM), WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        cls._add_run(p, title, bold=True, italic=italic_title, size=11)
        if date_str:
            p.add_run("\t")
            cls._add_run(p, date_str, size=10)
        return p

    @classmethod
    def _subline(cls, doc, text, *, italic=True, color="444444"):
        """Italic subline (company / tech stack)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        cls._add_run(p, text, italic=italic, size=10.5, color=color)
        return p

    @classmethod
    def _bullets(cls, doc, items):
        """Add bullet list using Word's List Bullet style."""
        added = False
        for b in items:
            if not b:
                continue
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            cls._add_run(p, b, size=10.5)
            added = True
        if added:
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    @staticmethod
    def _add_hyperlink(paragraph, text, url):
        """Append a blue underlined clickable hyperlink run to a paragraph."""
        r_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/"
            "relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "1155CC")
        rPr.append(color_el)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        rPr.append(fonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")  # 10pt = 20 half-points
        rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    @classmethod
    def _header_block(cls, doc, profile):
        """Name + contact line (with clickable links) + hr."""
        name = profile.get("name", "") or ""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        cls._add_run(p, name, bold=True, size=18)

        contact = profile.get("contact", {}) or {}
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)

        def _plain(text):
            r = p2.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        def _sep():
            _plain(" | ")

        plain_fields = [
            contact.get("email", ""),
            contact.get("phone", ""),
        ]
        link_fields = {
            "linkedin": contact.get("linkedin", ""),
            "github":   contact.get("github", ""),
            "website":  contact.get("website", ""),
        }

        # Emit non-empty plain fields first
        parts_plain = [f for f in plain_fields if f]
        for i, val in enumerate(parts_plain):
            _plain(val)
            if i < len(parts_plain) - 1 or any(link_fields.values()):
                _sep()

        # Emit link fields as proper hyperlinks
        link_items = [(v, v) for v in link_fields.values() if v]
        for i, (display, raw_url) in enumerate(link_items):
            url = raw_url if raw_url.startswith("http") else f"https://{raw_url}"
            cls._add_hyperlink(p2, display, url)
            if i < len(link_items) - 1:
                _sep()

        cls._set_bottom_border(p2)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(6)

    @classmethod
    def _skills_table(cls, doc, skills):
        """Two-column borderless table of skill categories."""
        ordered = ALWAYS_SHOW_SKILLS + CONDITIONAL_SKILLS
        rows = [(SKILL_LABELS[key], skills[key])
                for key in ordered
                if skills.get(key)]
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=2)
        table.autofit = False
        cls._clear_table_borders(table)
        # Force fixed table layout
        tblPr = table._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)

        left_w = 3.5
        right_w = PAGE_W_CM - 2 * MARGIN_CM - left_w
        for row_i, (label, values) in enumerate(rows):
            left = table.cell(row_i, 0)
            right = table.cell(row_i, 1)
            cls._set_cell_width(left, left_w)
            cls._set_cell_width(right, right_w)
            lp = left.paragraphs[0]
            lp.paragraph_format.space_before = Pt(1)
            lp.paragraph_format.space_after = Pt(1)
            cls._add_run(lp, f"{label}:", bold=True, size=10.5)
            rp = right.paragraphs[0]
            rp.paragraph_format.space_before = Pt(1)
            rp.paragraph_format.space_after = Pt(1)
            cls._add_run(rp, ", ".join(values), size=10.5)

    @staticmethod
    def _remove_compat_mode(doc):
        """Strip the w:compat block so Word opens in edit mode, not compat mode."""
        settings = doc.settings.element
        compat = settings.find(qn("w:compat"))
        if compat is not None:
            settings.remove(compat)
        doc.core_properties.revision = 1

    @classmethod
    def build(cls, profile, cv_data, output_path):
        """Build the full CV document."""
        doc = Document()
        cls._remove_compat_mode(doc)
        cls._set_defaults(doc)
        cls._set_page(doc)

        name = cv_data.get("name") or profile.get("name", "")
        contact = cv_data.get("contact") or profile.get("contact", {})
        cls._header_block(doc, {"name": name, "contact": contact})

        summary = cv_data.get("summary") or profile.get("summary")
        if summary:
            cls._section_header(doc, "Summary")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            cls._add_run(p, summary, size=10.5)

        experience = cv_data.get("experience") or []
        if experience:
            cls._section_header(doc, "Experience")
            for role in experience:
                dates = cls._fmt_range(
                    role.get("start_date") or role.get("start"),
                    role.get("end_date") or role.get("end"),
                    role.get("current"))
                cls._title_with_date(doc, role.get("title", ""), dates)
                comp = role.get("company", "")
                loc = role.get("location", "")
                sub = comp + (f" — {loc}" if loc else "")
                if sub.strip():
                    cls._subline(doc, sub)
                bullets = (role.get("bullets")
                           or role.get("responsibilities")
                           or role.get("achievements") or [])
                if not bullets:
                    print(f"[warn] empty bullets for role: {role.get('title')}")
                cls._bullets(doc, bullets)

        projects = cv_data.get("projects") or []
        if projects:
            cls._section_header(doc, "Projects")
            for proj in projects:
                date_str = cls._fmt_range(
                    proj.get("start_date"), proj.get("end_date"),
                    None) or str(proj.get("year") or "")
                proj_name = proj.get("name", "")
                proj_link = (proj.get("link") or "").strip()
                # Title line: name left, date right
                title_p = cls._title_with_date(doc, proj_name, date_str)
                # Append clickable link after the name if present
                if proj_link and proj_link.lower() != "github":
                    url = proj_link if proj_link.startswith("http") \
                        else f"https://{proj_link}"
                    title_p.add_run("  ")
                    cls._add_hyperlink(title_p, proj_link, url)
                elif proj_link.lower() == "github":
                    # "GitHub" label without a real URL — render as plain text
                    cls._add_run(title_p, "  GitHub", size=10, color="555555")
                tech = proj.get("technologies") or proj.get("tech") or []
                if isinstance(tech, list):
                    tech = ", ".join(tech)
                if tech:
                    cls._subline(doc, tech)
                bullets = proj.get("highlights") or proj.get("bullets") or []
                if not bullets:
                    desc = proj.get("description") or ""
                    if desc:
                        bullets = [s.strip() for s in re.split(r"\.\s+", desc)
                                   if s.strip()]
                if not bullets:
                    print(f"[warn] empty bullets for project: "
                          f"{proj.get('name')}")
                cls._bullets(doc, bullets)

        cv_skills = cv_data.get("skills") or {}
        profile_skills = profile.get("skills") or {}
        final_skills = {}
        for key in ALWAYS_SHOW_SKILLS:
            vals = cv_skills.get(key) or profile_skills.get(key) or []
            if vals:
                final_skills[key] = vals
        for key in CONDITIONAL_SKILLS:
            vals = cv_skills.get(key) or []
            if vals:
                final_skills[key] = vals
        if final_skills:
            cls._section_header(doc, "Technical Skills")
            cls._skills_table(doc, final_skills)

        volunteering = cv_data.get("volunteering") or []
        if volunteering:
            cls._section_header(doc, "Volunteering & Leadership")
            for v in volunteering:
                dates = cls._fmt_range(v.get("start_date"),
                                       v.get("end_date"), None) \
                        or v.get("period", "")
                cls._title_with_date(doc, v.get("role", ""), dates)
                org = v.get("organization") or v.get("org") or ""
                if org:
                    cls._subline(doc, org)
                bullets = v.get("bullets") or v.get("responsibilities") or []
                if not bullets:
                    desc = v.get("description") or ""
                    if desc:
                        bullets = [s.strip() for s in re.split(r"\.\s+", desc)
                                   if s.strip()]
                if not bullets:
                    print(f"[warn] empty bullets for volunteering: "
                          f"{v.get('role')}")
                cls._bullets(doc, bullets)

        achievements = cv_data.get("achievements") or []
        if achievements:
            cls._section_header(doc, "Achievements")
            cls._bullets(doc, achievements)

        education = cv_data.get("education") or profile.get("education") or []
        if education:
            cls._section_header(doc, "Education")
            for ed in education:
                dates = cls._fmt_range(ed.get("start_date"),
                                       ed.get("end_date"), None) \
                        or str(ed.get("year") or "")
                cls._title_with_date(doc, ed.get("degree", ""), dates)
                inst = ed.get("institution", "")
                if inst:
                    cls._subline(doc, inst)
                details = ed.get("details")
                if details:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(4)
                    cls._add_run(p, details, size=10, color="444444")

        doc.save(output_path)

    @staticmethod
    def _fmt_range(start, end, current):
        """Format a date range. Accepts start_date/end_date or start/end."""
        if current:
            end = "Present"
        start = (start or "").strip() if isinstance(start, str) else start
        end = (end or "").strip() if isinstance(end, str) else end
        if start and end:
            return f"{start} — {end}"
        return start or end or ""


# ─── Unified Assess+Generate Worker (one call instead of two) ────────────

UNIFIED_SYSTEM = (
    "You are a CV writer. Generate a tailored CV using ONLY the "
    "candidate's profile data. Do not invent skills, projects, metrics, "
    "or achievements.\n\n"
    "Return JSON only, no markdown:\n"
    "{\"fit\":{\"fit\":\"green|yellow|red\",\"score\":0-100,"
    "\"summary\":\"1 sentence\",\"strengths\":[],\"gaps\":[],"
    "\"hard_gaps\":[]},"
    "\"cv\":{"
    "\"summary\":\"max 3 sentences tailored to JD\","
    "\"experience\":[{\"title\":\"\",\"company\":\"\","
    "\"start_date\":\"\",\"end_date\":\"\",\"bullets\":[]}],"
    "\"projects\":[{\"name\":\"\",\"technologies\":[],\"bullets\":[]}],"
    "\"skills\":{JD-relevant categories only; items verbatim — see SKILLS SECTION RULES},"
    "\"volunteering\":[{\"role\":\"\",\"organization\":\"\","
    "\"start_date\":\"\",\"end_date\":\"\",\"bullets\":[]}],"
    "\"achievements\":[verbatim from profile],"
    "\"education\":[verbatim from profile]},"
    "\"hard_gap\":\"one sentence if hard gap else empty\"}\n\n"
    "Rules:\n"
    "- Section order: summary, experience, projects, skills, "
    "volunteering, achievements, education\n"
    "- Include ONLY the 3 most relevant projects for the JD\n"
    "- Max 2 bullets per experience entry (most JD-relevant first)\n"
    "- Max 2 bullets per volunteering entry\n"
    "- Max 3 bullets per project entry (most JD-relevant first)\n"
    "- Use verbatim skills/achievements from profile — never modify\n"
    "- Mirror JD keywords using real profile content only\n"
    "- Fit: green=70+, yellow=40-69, red=0-39. Be strict.\n"
    "- SKILLS SECTION RULES:\n"
    "  Always include these core categories (verbatim from profile): "
    "languages, frontend, backend, databases, cloud, ai_integrations, "
    "third_party_apis\n"
    "  Include these ONLY if the JD explicitly mentions or clearly implies "
    "them (e.g. 'Excel'/'spreadsheets' → productivity_tools; "
    "'Figma'/'design' → design_collaboration): "
    "erp, desktop_gui, productivity_tools, design_collaboration, "
    "analytics_tools, dev_tools, soft_skills, languages_spoken\n"
    "  Never include a conditional category just to pad the CV"
)


class UnifiedWorker(QThread):
    """Single-call fit assessment + CV generation."""
    progress = pyqtSignal(int)
    result = pyqtSignal(dict, dict, str)  # fit, cv, hard_gap
    error = pyqtSignal(str)

    def __init__(self, client, profile, company, title, jd):
        """Store inputs."""
        super().__init__()
        self.client = client
        self.profile = profile
        self.company = company
        self.title = title
        self.jd = jd

    def run(self):
        """Execute the combined assessment+generation call."""
        compact = compress_profile(self.profile)
        cached_user = f"CANDIDATE PROFILE:\n{compact}"
        fresh_user = (f"COMPANY: {self.company}\nTITLE: {self.title}\n"
                      f"JOB DESCRIPTION:\n{self.jd}\n\n"
                      f"Generate the tailored CV.")
        self.progress.emit(15)
        try:
            text = claude_call_cached(self.client, UNIFIED_SYSTEM,
                                      cached_user, fresh_user,
                                      MAX_TOKENS_UNIFIED)
            self.progress.emit(80)
            data = self._parse(text)
            fit = data.get("fit") or {}
            cv = validate_cv_output(data.get("cv") or {}, self.profile)
            hard_gap = data.get("hard_gap", "") or ""
            self.progress.emit(100)
            self.result.emit(fit, cv, hard_gap)
        except AuthenticationError:
            self.error.emit("Invalid API key — check Settings")
        except APIConnectionError:
            self.error.emit("Connection failed — check internet")
        except Exception as e:
            self.error.emit(str(e))

    def _parse(self, text):
        """Parse JSON; on failure retry once with an explicit format reminder."""
        body, _ = strip_hard_gap(text)
        try:
            return parse_json_response(body)
        except Exception:
            pass
        # Retry with a hard format instruction
        retry_suffix = (
            "\nIMPORTANT: Return ONLY a raw JSON object. No markdown, "
            "no backticks, no explanation. Start your response with { "
            "and end with }"
        )
        cached_user = f"CANDIDATE PROFILE:\n{compress_profile(self.profile)}"
        fresh_user = (f"COMPANY: {self.company}\nTITLE: {self.title}\n"
                      f"JOB DESCRIPTION:\n{self.jd}\n\n"
                      f"Generate the tailored CV.")
        text2 = claude_call_cached(
            self.client,
            UNIFIED_SYSTEM + retry_suffix,
            cached_user, fresh_user,
            MAX_TOKENS_UNIFIED,
        )
        body2, _ = strip_hard_gap(text2)
        try:
            return parse_json_response(body2)
        except Exception as exc:
            raise ValueError(
                "CV generation failed — Claude returned unexpected format. "
                "Check console for details. Try again."
            ) from exc


class ProfileBuildWorker(QThread):
    """Build or merge a profile in the background."""
    progress = pyqtSignal(int)
    status = pyqtSignal(str)
    done = pyqtSignal(dict)
    error = pyqtSignal(str)

    def __init__(self, client, files, mode, existing=None):
        """Store inputs."""
        super().__init__()
        self.client = client
        self.files = files
        self.mode = mode
        self.existing = existing

    def run(self):
        """Extract texts and call ProfileManager."""
        try:
            texts = []
            total = max(1, len(self.files))
            for i, f in enumerate(self.files):
                self.status.emit(f"Reading {os.path.basename(f)}…")
                texts.append(extract_text_from_file(f))
                self.progress.emit(int((i + 1) / total * 50))
            self.status.emit("Calling Claude…")
            pm = ProfileManager(self.client)
            if self.mode == "new":
                profile = pm.build_new(texts)
            else:
                profile = pm.merge(self.existing or PROFILE_SCHEMA, texts)
            self.progress.emit(100)
            self.done.emit(profile)
        except AuthenticationError:
            self.error.emit("Invalid API key — check Settings")
        except APIConnectionError:
            self.error.emit("Connection failed — check internet")
        except Exception as e:
            self.error.emit(str(e))


# ─── Tabs ────────────────────────────────────────────────────────────────

class SettingsTab(QWidget):
    """API key and output folder configuration."""
    config_changed = pyqtSignal()

    def __init__(self, cfg):
        """Build UI."""
        super().__init__()
        self.cfg = cfg
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)

        layout.addWidget(QLabel("Anthropic API Key"))
        key_row = QHBoxLayout()
        stored_key = cfg.get("api_key", "")
        display_key = stored_key if stored_key.startswith("sk-ant-") else ""
        self.key_edit = QLineEdit(display_key)
        self.key_edit.setEchoMode(QLineEdit.Password)
        self.show_btn = QPushButton("Show")
        self.show_btn.setCheckable(True)
        self.show_btn.toggled.connect(self._toggle_show)
        key_row.addWidget(self.key_edit)
        key_row.addWidget(self.show_btn)
        layout.addLayout(key_row)

        layout.addSpacing(10)
        layout.addWidget(QLabel("Output Folder"))
        folder_row = QHBoxLayout()
        self.folder_edit = QLineEdit(cfg.get("output_folder", DEFAULT_OUTPUT))
        browse = QPushButton("Browse…")
        browse.clicked.connect(self._browse)
        folder_row.addWidget(self.folder_edit)
        folder_row.addWidget(browse)
        layout.addLayout(folder_row)

        layout.addSpacing(15)
        save_row = QHBoxLayout()
        self.save_btn = QPushButton("Save")
        self.save_btn.clicked.connect(self._save)
        self.saved_lbl = QLabel("")
        save_row.addWidget(self.save_btn)
        save_row.addWidget(self.saved_lbl)
        save_row.addStretch()
        layout.addLayout(save_row)

        layout.addStretch()
        ver = QLabel(f"CV Tailor v{APP_VERSION}")
        ver.setStyleSheet("color:#8b949e;")
        layout.addWidget(ver)

    def _toggle_show(self, checked):
        """Toggle API key visibility."""
        self.key_edit.setEchoMode(
            QLineEdit.Normal if checked else QLineEdit.Password)
        self.show_btn.setText("Hide" if checked else "Show")

    def _browse(self):
        """Pick output folder."""
        d = QFileDialog.getExistingDirectory(self, "Select Output Folder",
                                             self.folder_edit.text())
        if d:
            self.folder_edit.setText(d)

    def _save(self):
        """Save config to disk, rejecting keys with wrong prefix."""
        key = self.key_edit.text().strip()
        if key and not key.startswith("sk-ant-"):
            self.saved_lbl.setText(
                "Invalid API key format — should start with sk-ant-")
            self.saved_lbl.setStyleSheet(f"color:{RED};")
            return
        self.cfg["api_key"] = key
        self.cfg["output_folder"] = self.folder_edit.text().strip() \
            or DEFAULT_OUTPUT
        save_config(self.cfg)
        os.makedirs(self.cfg["output_folder"], exist_ok=True)
        self.saved_lbl.setText("Saved ✓")
        self.saved_lbl.setStyleSheet(f"color:{GREEN};")
        self.config_changed.emit()


class ProfileTab(QWidget):
    """Build/merge master profile, view summary."""
    profile_changed = pyqtSignal()

    def __init__(self, get_client, get_profile, set_profile):
        """Build UI."""
        super().__init__()
        self.get_client = get_client
        self.get_profile = get_profile
        self.set_profile = set_profile
        self.files = []
        self.worker = None

        outer = QHBoxLayout(self)
        outer.setContentsMargins(15, 15, 15, 15)

        left = QGroupBox("Build / Update Profile")
        lv = QVBoxLayout(left)
        pick = QPushButton("Add Files (PDF, DOCX, TXT)…")
        pick.clicked.connect(self._pick_files)
        lv.addWidget(pick)
        self.file_list = QListWidget()
        lv.addWidget(self.file_list)
        btn_row = QHBoxLayout()
        self.build_btn = QPushButton("Build New Profile")
        self.merge_btn = QPushButton("Merge Into Profile")
        self.build_btn.clicked.connect(lambda: self._run("new"))
        self.merge_btn.clicked.connect(lambda: self._run("merge"))
        btn_row.addWidget(self.build_btn)
        btn_row.addWidget(self.merge_btn)
        lv.addLayout(btn_row)
        self.progress = QProgressBar()
        lv.addWidget(self.progress)
        self.log = QTextEdit()
        self.log.setReadOnly(True)
        self.log.setMaximumHeight(150)
        lv.addWidget(self.log)
        view_raw = QPushButton("View Raw JSON")
        view_raw.clicked.connect(self._view_raw)
        lv.addWidget(view_raw)

        right = QGroupBox("Current Profile Summary")
        rv = QVBoxLayout(right)
        self.summary = QTextEdit()
        self.summary.setReadOnly(True)
        rv.addWidget(self.summary)
        clear = QPushButton("Clear Profile")
        clear.clicked.connect(self._clear)
        rv.addWidget(clear)

        outer.addWidget(left, 1)
        outer.addWidget(right, 1)
        self.refresh_summary()

    def _pick_files(self):
        """Add files to the queue."""
        files, _ = QFileDialog.getOpenFileNames(
            self, "Select CV Documents", "",
            "Documents (*.pdf *.docx *.txt)")
        for f in files:
            if f not in self.files:
                self.files.append(f)
                item = QListWidgetItem(os.path.basename(f))
                item.setData(Qt.UserRole, f)
                self.file_list.addItem(item)

    def _run(self, mode):
        """Kick off build/merge worker."""
        client = self.get_client()
        if not client:
            QMessageBox.warning(self, "No API Key",
                                "Set your API key in Settings first.")
            return
        if not self.files:
            QMessageBox.warning(self, "No Files", "Add at least one file.")
            return
        self.build_btn.setEnabled(False)
        self.merge_btn.setEnabled(False)
        self.log.append(f"\n— {mode.upper()} —")
        self.progress.setValue(0)
        existing = self.get_profile() if mode == "merge" else None
        self.worker = ProfileBuildWorker(client, self.files, mode, existing)
        self.worker.status.connect(lambda s: self.log.append(s))
        self.worker.progress.connect(self.progress.setValue)
        self.worker.done.connect(self._on_done)
        self.worker.error.connect(self._on_error)
        self.worker.start()

    def _on_done(self, profile):
        """Persist profile and update UI."""
        save_profile(profile)
        self.set_profile(profile)
        roles = len(profile.get("experience") or [])
        projs = len(profile.get("projects") or [])
        skills = sum(len(profile.get("skills", {}).get(k, []))
                     for k, _ in SKILL_CATEGORIES)
        self.log.append(f"Done — {roles} roles, {skills} skills, "
                        f"{projs} projects.")
        self.build_btn.setEnabled(True)
        self.merge_btn.setEnabled(True)
        self.refresh_summary()
        self.profile_changed.emit()

    def _on_error(self, msg):
        """Show error in log."""
        self.log.append(f"ERROR: {msg}")
        self.build_btn.setEnabled(True)
        self.merge_btn.setEnabled(True)

    def _view_raw(self):
        """Open master_profile.json in Notepad."""
        if os.path.exists(PROFILE_FILE):
            subprocess.Popen(["notepad.exe", PROFILE_FILE])
        else:
            QMessageBox.information(self, "No Profile", "No profile yet.")

    def _clear(self):
        """Clear profile after confirmation."""
        r = QMessageBox.question(self, "Clear Profile",
                                 "Delete the master profile?")
        if r == QMessageBox.Yes:
            if os.path.exists(PROFILE_FILE):
                os.remove(PROFILE_FILE)
            self.set_profile(None)
            self.refresh_summary()
            self.profile_changed.emit()

    def refresh_summary(self):
        """Refresh the right-hand summary panel."""
        p = self.get_profile()
        if not p:
            self.summary.setPlainText("No profile loaded.")
            return
        contact = p.get("contact", {}) or {}
        lines = [f"Name: {p.get('name', '')}",
                 f"Email: {contact.get('email', '')}",
                 f"Phone: {contact.get('phone', '')}",
                 f"LinkedIn: {contact.get('linkedin', '')}",
                 f"GitHub: {contact.get('github', '')}", ""]
        roles = p.get("experience") or []
        lines.append(f"Roles ({len(roles)}):")
        for r in roles:
            lines.append(f"  • {r.get('title', '')} @ {r.get('company', '')}")
        projs = p.get("projects") or []
        lines.append(f"\nProjects ({len(projs)}):")
        for pr in projs:
            lines.append(f"  • {pr.get('name', '')}")
        lines.append("\nSkills:")
        for k, label in SKILL_CATEGORIES:
            vals = p.get("skills", {}).get(k) or []
            if vals:
                lines.append(f"  {label}: {', '.join(vals)}")
        self.summary.setPlainText("\n".join(lines))


class SingleJobTab(QWidget):
    """Paste one JD, assess + generate in a single call."""

    def __init__(self, get_client, get_profile, get_output):
        """Build UI."""
        super().__init__()
        self.get_client = get_client
        self.get_profile = get_profile
        self.get_output = get_output
        self.worker = None
        self.output_path = None
        self._cached = None  # (fit, cv, hard_gap) pending RED decision

        v = QVBoxLayout(self)
        v.setContentsMargins(15, 15, 15, 15)
        form = QFormLayout()
        self.company = QLineEdit()
        self.title = QLineEdit()
        form.addRow("Company:", self.company)
        form.addRow("Job Title:", self.title)
        v.addLayout(form)

        v.addWidget(QLabel("Paste the full job description here:"))
        self.jd = QTextEdit()
        self.jd.textChanged.connect(self._update_state)
        v.addWidget(self.jd)
        self.charlbl = QLabel("0 characters")
        v.addWidget(self.charlbl)

        self.company.textChanged.connect(self._update_state)
        self.title.textChanged.connect(self._update_state)

        self.go_btn = QPushButton("Assess & Generate CV")
        self.go_btn.clicked.connect(self._start)
        v.addWidget(self.go_btn)

        self.phase_lbl = QLabel("")
        v.addWidget(self.phase_lbl)
        self.progress = QProgressBar()
        v.addWidget(self.progress)

        self.fit_box = QLabel("")
        self.fit_box.setWordWrap(True)
        self.fit_box.setVisible(False)
        v.addWidget(self.fit_box)

        self.red_row = QWidget()
        rh = QHBoxLayout(self.red_row)
        self.anyway_btn = QPushButton("Generate Anyway")
        self.skip_btn = QPushButton("Skip This Job")
        self.anyway_btn.clicked.connect(self._red_generate)
        self.skip_btn.clicked.connect(self._red_skip)
        rh.addWidget(self.anyway_btn)
        rh.addWidget(self.skip_btn)
        rh.addStretch()
        self.red_row.setVisible(False)
        v.addWidget(self.red_row)

        self.file_lbl = QLabel("")
        v.addWidget(self.file_lbl)
        file_row = QHBoxLayout()
        self.open_file_btn = QPushButton("Open File")
        self.open_folder_btn = QPushButton("Open Output Folder")
        self.open_file_btn.clicked.connect(
            lambda: self.output_path and open_file_native(self.output_path))
        self.open_folder_btn.clicked.connect(
            lambda: open_file_native(self.get_output()))
        self.open_file_btn.setVisible(False)
        self.open_folder_btn.setVisible(False)
        file_row.addWidget(self.open_file_btn)
        file_row.addWidget(self.open_folder_btn)
        file_row.addStretch()
        v.addLayout(file_row)

        self.gap_banner = QLabel("")
        self.gap_banner.setWordWrap(True)
        self.gap_banner.setVisible(False)
        v.addWidget(self.gap_banner)

        self._update_state()

    def _update_state(self):
        """Enable Go button only if all fields filled."""
        self.charlbl.setText(f"{len(self.jd.toPlainText())} characters")
        ok = (self.company.text().strip() and self.title.text().strip()
              and self.jd.toPlainText().strip())
        self.go_btn.setEnabled(bool(ok))

    def _start(self):
        """Kick off unified assess+generate call."""
        client = self.get_client()
        profile = self.get_profile()
        if not client:
            QMessageBox.warning(self, "No API Key", "Set your API key.")
            return
        if not profile:
            QMessageBox.warning(self, "No Profile",
                                "Build your master profile first.")
            return
        self.fit_box.setVisible(False)
        self.red_row.setVisible(False)
        self.file_lbl.setText("")
        self.gap_banner.setVisible(False)
        self.open_file_btn.setVisible(False)
        self.open_folder_btn.setVisible(False)
        self.go_btn.setEnabled(False)
        self.phase_lbl.setText("Assessing fit & drafting CV…")
        self.progress.setValue(0)
        self.worker = UnifiedWorker(client, profile, self.company.text(),
                                    self.title.text(), self.jd.toPlainText())
        self.worker.progress.connect(self.progress.setValue)
        self.worker.result.connect(self._on_result)
        self.worker.error.connect(self._on_error)
        self.worker.start()

    def _on_result(self, fit, cv, hard_gap):
        """Handle unified result: render fit + gate on RED."""
        level = (fit.get("fit") or "yellow").lower()
        summary = fit.get("summary", "")
        gaps = fit.get("gaps", []) or []
        hard = fit.get("hard_gaps", []) or []
        self._cached = (fit, cv, hard_gap)

        if level == "green":
            self._show_fit(GREEN, f"Strong fit — generating your CV\n{summary}")
            self._write_cv()
        elif level == "yellow":
            g = "; ".join(gaps) if gaps else "minor"
            self._show_fit(YELLOW,
                           f"Partial fit — {g} — generating anyway\n{summary}")
            self._write_cv()
        else:
            reason = summary + ("\nHard gaps: " + "; ".join(hard)
                                if hard else "")
            self._show_fit(RED, f"Poor fit — {reason}")
            self.red_row.setVisible(True)
            self.phase_lbl.setText("Waiting for your decision…")

    def _show_fit(self, color, text):
        """Display the fit box."""
        self.fit_box.setText(text)
        self.fit_box.setStyleSheet(
            f"background:{color};color:white;padding:10px;border-radius:4px;")
        self.fit_box.setVisible(True)

    def _red_generate(self):
        """User overrode RED — write CV from cache."""
        self.red_row.setVisible(False)
        self._write_cv()

    def _red_skip(self):
        """User skipped RED."""
        self.red_row.setVisible(False)
        self.phase_lbl.setText("Skipped.")
        self._cached = None
        self.go_btn.setEnabled(True)

    def _write_cv(self):
        """Write cached CV data to DOCX."""
        if not self._cached:
            self.go_btn.setEnabled(True)
            return
        _, cv, hard_gap = self._cached
        try:
            profile = self.get_profile()
            self.output_path = build_output_path(self.get_output(),
                                                 self.company.text(),
                                                 self.title.text())
            DocxBuilder.build(profile, cv, self.output_path)
            self.phase_lbl.setText("Done.")
            self.file_lbl.setText(
                f"Generated: {os.path.basename(self.output_path)}")
            self.open_file_btn.setVisible(True)
            self.open_folder_btn.setVisible(True)
            if hard_gap:
                self.gap_banner.setText(f"⚠ Hard gap: {hard_gap}")
                self.gap_banner.setStyleSheet(
                    f"background:{YELLOW};color:white;padding:8px;"
                    "border-radius:4px;")
                self.gap_banner.setVisible(True)
        except Exception as e:
            self._on_error(f"DOCX build failed: {e}")
        self.go_btn.setEnabled(True)

    def _on_error(self, msg):
        """Show error state."""
        self.progress.setValue(0)
        self.phase_lbl.setText(f"Error: {msg}")
        self.go_btn.setEnabled(True)


# ─── Bulk Tab ────────────────────────────────────────────────────────────

def parse_bulk_input(text):
    """Parse bulk input into list of (company, title, jd) tuples."""
    jobs = []
    blocks = re.split(r"\n\s*---\s*\n", text.strip())
    for block in blocks:
        if not block.strip():
            continue
        comp = re.search(r"COMPANY\s*:\s*(.+)", block, re.I)
        title = re.search(r"TITLE\s*:\s*(.+)", block, re.I)
        jd_m = re.search(r"JD\s*:\s*\n?(.+)", block, re.I | re.S)
        if comp and title and jd_m:
            jobs.append((comp.group(1).strip(), title.group(1).strip(),
                         jd_m.group(1).strip()))
    return jobs


class BulkRunner(QThread):
    """Sequentially assess+generate with RED pause support. One API call
    per job via UNIFIED_SYSTEM."""
    row_update = pyqtSignal(int, str, str)
    waiting_for_decision = pyqtSignal(int, str, list)
    done = pyqtSignal()

    def __init__(self, client, profile, output_folder, jobs):
        """Store inputs."""
        super().__init__()
        self.client = client
        self.profile = profile
        self.output = output_folder
        self.jobs = jobs
        self._stop = False
        self._decision_event = threading.Event()
        self._decision = None
        self.results = []

    def stop(self):
        """Request stop — unblocks any pending decision wait."""
        self._stop = True
        self._decision_event.set()

    def submit_decision(self, decision):
        """Called from GUI with 'generate' or 'skip'."""
        self._decision = decision
        self._decision_event.set()

    def run(self):
        """Main loop."""
        for i, (company, title, jd) in enumerate(self.jobs):
            if self._stop:
                break
            row_result = {"Company": company, "Title": title,
                          "Fit": "", "Fit Score": "", "Fit Summary": "",
                          "Strengths": "", "Gaps": "", "Hard Gaps": "",
                          "Status": "", "Gap Note": "", "Filename": "",
                          "Date": datetime.now().strftime("%Y-%m-%d")}
            self.row_update.emit(i, "Status", "Assessing & drafting")
            fit, cv, hard_gap = self._unified(company, title, jd)
            if fit is None:
                self.row_update.emit(i, "Status", "✗ Error")
                row_result["Status"] = "Error"
                self.results.append(row_result)
                time.sleep(BULK_DELAY_SEC)
                continue

            level = (fit.get("fit") or "yellow").lower()
            icon = {"green": "🟢 Strong", "yellow": "🟡 Partial",
                    "red": "🔴 Poor"}.get(level, "🟡 Partial")
            self.row_update.emit(i, "Fit", icon)
            row_result["Fit"] = icon
            row_result["Fit Score"] = str(fit.get("score", ""))
            row_result["Fit Summary"] = fit.get("summary", "")
            row_result["Strengths"] = "; ".join(fit.get("strengths", []))
            row_result["Gaps"] = "; ".join(fit.get("gaps", []))
            row_result["Hard Gaps"] = "; ".join(fit.get("hard_gaps", []))

            if level == "red":
                self.row_update.emit(i, "Status", "⚠ Poor Fit — waiting...")
                self._decision_event.clear()
                self._decision = None
                self.waiting_for_decision.emit(
                    i, fit.get("summary", ""), fit.get("hard_gaps", []))
                self._decision_event.wait()
                if self._stop:
                    row_result["Status"] = "Skipped"
                    self.results.append(row_result)
                    break
                if self._decision != "generate":
                    self.row_update.emit(i, "Status", "✗ Skipped")
                    row_result["Status"] = "Skipped"
                    self.results.append(row_result)
                    time.sleep(BULK_DELAY_SEC)
                    continue

            self.row_update.emit(i, "Status", "Writing DOCX")
            try:
                path = build_output_path(self.output, company, title)
                DocxBuilder.build(self.profile, cv or {}, path)
                self.row_update.emit(i, "Status", "✓ Done")
                self.row_update.emit(i, "Filename", os.path.basename(path))
                row_result["Status"] = "Done"
                row_result["Filename"] = os.path.basename(path)
                if hard_gap:
                    self.row_update.emit(i, "Gap", f"⚠ {hard_gap}")
                    row_result["Gap Note"] = hard_gap
            except Exception as e:
                self.row_update.emit(i, "Status", f"✗ Error: {e}")
                row_result["Status"] = f"Error: {e}"
            self.results.append(row_result)
            time.sleep(BULK_DELAY_SEC)
        self.done.emit()

    def _unified(self, company, title, jd):
        """Single combined fit+CV call. Returns (fit, cv, hard_gap) or
        (None, None, '') on error."""
        compact = compress_profile(self.profile)
        cached_user = f"CANDIDATE PROFILE:\n{compact}"
        fresh_user = (f"COMPANY: {company}\nTITLE: {title}\n"
                      f"JOB DESCRIPTION:\n{jd}\n\n"
                      f"Generate the tailored CV.")
        try:
            text = claude_call_cached(self.client, UNIFIED_SYSTEM,
                                      cached_user, fresh_user,
                                      MAX_TOKENS_UNIFIED)
            body, _ = strip_hard_gap(text)
            data = parse_json_response(body)
            return (data.get("fit") or {},
                    validate_cv_output(data.get("cv") or {}, self.profile),
                    data.get("hard_gap", "") or "")
        except AuthenticationError:
            return None, None, ""
        except APIConnectionError:
            return None, None, ""
        except Exception:
            # Parse retry with explicit format reminder
            retry_suffix = (
                "\nIMPORTANT: Return ONLY a raw JSON object. No markdown, "
                "no backticks, no explanation. Start your response with { "
                "and end with }"
            )
            try:
                text = claude_call_cached(
                    self.client, UNIFIED_SYSTEM + retry_suffix,
                    cached_user, fresh_user, MAX_TOKENS_UNIFIED)
                body, _ = strip_hard_gap(text)
                data = parse_json_response(body)
                return (data.get("fit") or {},
                        validate_cv_output(data.get("cv") or {}, self.profile),
                        data.get("hard_gap", "") or "")
            except Exception:
                return None, None, ""


class BulkTab(QWidget):
    """Bulk-process many JDs at once."""

    COLS = ["#", "Company", "Title", "Fit", "Status", "Gap", "Filename"]

    def __init__(self, get_client, get_profile, get_output):
        """Build UI."""
        super().__init__()
        self.get_client = get_client
        self.get_profile = get_profile
        self.get_output = get_output
        self.runner = None
        self.pending_rows = {}

        v = QVBoxLayout(self)
        v.setContentsMargins(15, 15, 15, 15)

        self.instructions = QTextEdit()
        self.instructions.setReadOnly(True)
        self.instructions.setMaximumHeight(140)
        self.instructions.setPlainText(
            "Format (separate jobs with ---):\n\n"
            "COMPANY: Google\nTITLE: Software Engineer\nJD:\n"
            "[full job description here]\n---\n"
            "COMPANY: Meta\nTITLE: Full Stack Developer\nJD:\n"
            "[full job description here]\n---")
        v.addWidget(self.instructions)

        self.input = QTextEdit()
        self.input.textChanged.connect(self._update_count)
        v.addWidget(self.input)

        self.count_lbl = QLabel("0 jobs detected")
        v.addWidget(self.count_lbl)

        ctrl = QHBoxLayout()
        self.start_btn = QPushButton("Start")
        self.stop_btn = QPushButton("Stop")
        self.stop_btn.setEnabled(False)
        self.start_btn.clicked.connect(self._start)
        self.stop_btn.clicked.connect(self._stop)
        ctrl.addWidget(self.start_btn)
        ctrl.addWidget(self.stop_btn)
        ctrl.addStretch()
        v.addLayout(ctrl)

        self.progress = QProgressBar()
        v.addWidget(self.progress)

        self.table = QTableWidget(0, len(self.COLS))
        self.table.setHorizontalHeaderLabels(self.COLS)
        self.table.horizontalHeader().setSectionResizeMode(
            QHeaderView.Stretch)
        self.table.cellDoubleClicked.connect(self._open_row)
        v.addWidget(self.table)

        self.summary_lbl = QLabel("")
        v.addWidget(self.summary_lbl)

        bot = QHBoxLayout()
        open_out = QPushButton("Open Output Folder")
        open_out.clicked.connect(
            lambda: open_file_native(self.get_output()))
        export = QPushButton("Export Status CSV")
        export.clicked.connect(self._export_csv)
        bot.addWidget(open_out)
        bot.addWidget(export)
        bot.addStretch()
        v.addLayout(bot)

    def _update_count(self):
        """Refresh detected job count."""
        n = len(parse_bulk_input(self.input.toPlainText()))
        self.count_lbl.setText(f"{n} jobs detected")

    def _start(self):
        """Begin bulk run."""
        client = self.get_client()
        profile = self.get_profile()
        if not client or not profile:
            QMessageBox.warning(self, "Missing",
                                "Set API key and build profile first.")
            return
        jobs = parse_bulk_input(self.input.toPlainText())
        if not jobs:
            QMessageBox.warning(self, "No Jobs", "No jobs parsed.")
            return
        self.table.setRowCount(0)
        for i, (c, t, _) in enumerate(jobs):
            self.table.insertRow(i)
            self._set(i, 0, str(i + 1))
            self._set(i, 1, c)
            self._set(i, 2, t)
            self._set(i, 3, "")
            self._set(i, 4, "Queued")
            self._set(i, 5, "")
            self._set(i, 6, "")
        self.progress.setMaximum(len(jobs))
        self.progress.setValue(0)
        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.runner = BulkRunner(client, profile, self.get_output(), jobs)
        self.runner.row_update.connect(self._on_row)
        self.runner.waiting_for_decision.connect(self._on_red)
        self.runner.done.connect(self._on_done)
        self.runner.start()

    def _stop(self):
        """Request stop."""
        if self.runner:
            self.runner.stop()
        self.stop_btn.setEnabled(False)

    def _set(self, row, col, text):
        """Set a cell by index."""
        self.table.setItem(row, col, QTableWidgetItem(text))

    def _col_index(self, name):
        """Map column name to index."""
        return {"Status": 4, "Fit": 3, "Gap": 5, "Filename": 6}.get(name, 4)

    def _on_row(self, row, col_name, value):
        """Update a row cell."""
        col = self._col_index(col_name)
        item = self.table.item(row, col)
        if item:
            item.setText(value)
        if col_name == "Status":
            if value.startswith("✓") or value.startswith("✗"):
                self.progress.setValue(self.progress.value() + 1)
                if value.startswith("✗") and "Skipped" not in value:
                    self._highlight_row(row, "#3d1a1a")
            if "Poor Fit" in value:
                self._highlight_row(row, "#3d1a1a")

    def _highlight_row(self, row, color):
        """Paint an entire row."""
        for c in range(self.table.columnCount()):
            it = self.table.item(row, c)
            if it:
                it.setBackground(QColor(color))

    def _on_red(self, row, summary, hard_gaps):
        """Show inline Generate/Skip buttons in the row."""
        w = QWidget()
        h = QHBoxLayout(w)
        h.setContentsMargins(2, 2, 2, 2)
        gen = QPushButton("Generate Anyway")
        skp = QPushButton("Skip")
        gen.clicked.connect(lambda: self._decide(row, "generate"))
        skp.clicked.connect(lambda: self._decide(row, "skip"))
        h.addWidget(gen)
        h.addWidget(skp)
        self.table.setCellWidget(row, 4, w)
        self.pending_rows[row] = w

    def _decide(self, row, choice):
        """Deliver user's RED decision to the runner."""
        if row in self.pending_rows:
            self.table.removeCellWidget(row, 4)
            del self.pending_rows[row]
        if self.runner:
            self.runner.submit_decision(choice)

    def _on_done(self):
        """Finalise UI after runner completes."""
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        done = skipped = errors = gaps = 0
        for r in self.runner.results if self.runner else []:
            if r["Status"] == "Done":
                done += 1
            elif r["Status"] == "Skipped":
                skipped += 1
            elif r["Status"].startswith("Error"):
                errors += 1
            if r["Gap Note"]:
                gaps += 1
        self.summary_lbl.setText(
            f"{done} done · {skipped} skipped · {errors} errors · "
            f"{gaps} with gaps")

    def _open_row(self, row, _col):
        """Open the CV file when a Done row is double-clicked."""
        status_item = self.table.item(row, 4)
        fname_item = self.table.item(row, 6)
        if status_item and "Done" in status_item.text() and fname_item:
            path = os.path.join(self.get_output(), fname_item.text())
            if os.path.exists(path):
                open_file_native(path)

    def _export_csv(self):
        """Export results to CSV."""
        if not self.runner or not self.runner.results:
            QMessageBox.information(self, "No Data", "Nothing to export.")
            return
        date = datetime.now().strftime("%Y-%m-%d")
        default = os.path.join(self.get_output(),
                               f"cv_tailor_status_{date}.csv")
        path, _ = QFileDialog.getSaveFileName(self, "Export CSV", default,
                                              "CSV (*.csv)")
        if not path:
            return
        cols = ["Date", "Company", "Title", "Fit", "Fit Score",
                "Fit Summary", "Strengths", "Gaps", "Hard Gaps",
                "Status", "Gap Note", "Filename"]
        with open(path, "w", newline="", encoding="utf-8") as f:
            w = csv.DictWriter(f, fieldnames=cols)
            w.writeheader()
            for r in self.runner.results:
                w.writerow({c: r.get(c, "") for c in cols})
        QMessageBox.information(self, "Exported", f"Saved to {path}")


# ─── MainWindow ──────────────────────────────────────────────────────────

class MainWindow(QMainWindow):
    """Top-level window holding all tabs."""

    def __init__(self):
        """Build UI."""
        super().__init__()
        self.setWindowTitle("CV Tailor")
        self.resize(1150, 800)
        self.cfg = load_config()
        self.profile = load_profile()
        os.makedirs(self.cfg.get("output_folder", DEFAULT_OUTPUT),
                    exist_ok=True)
        self._client = None
        self._build_ui()
        self._apply_theme()
        self._refresh_banner()

    def _build_ui(self):
        """Construct tabs."""
        container = QWidget()
        v = QVBoxLayout(container)
        v.setContentsMargins(0, 0, 0, 0)
        self.banner = QLabel("")
        self.banner.setWordWrap(True)
        self.banner.setStyleSheet(
            f"background:{RED};color:white;padding:8px;")
        self.banner.setVisible(False)
        v.addWidget(self.banner)

        self.tabs = QTabWidget()
        self.settings_tab = SettingsTab(self.cfg)
        self.settings_tab.config_changed.connect(self._on_config)

        self.profile_tab = ProfileTab(
            self._get_client, lambda: self.profile, self._set_profile)
        self.profile_tab.profile_changed.connect(self._refresh_banner)

        self.single_tab = SingleJobTab(
            self._get_client, lambda: self.profile,
            lambda: self.cfg.get("output_folder", DEFAULT_OUTPUT))
        self.bulk_tab = BulkTab(
            self._get_client, lambda: self.profile,
            lambda: self.cfg.get("output_folder", DEFAULT_OUTPUT))

        self.tabs.addTab(self.settings_tab, "Settings")
        self.tabs.addTab(self.profile_tab, "Profile")
        self.tabs.addTab(self.single_tab, "Single Job")
        self.tabs.addTab(self.bulk_tab, "Bulk Jobs")
        v.addWidget(self.tabs)
        self.setCentralWidget(container)

    def _get_client(self):
        """Return or create the Anthropic client."""
        key = self.cfg.get("api_key", "").strip()
        if not key:
            return None
        if self._client is None:
            try:
                self._client = Anthropic(api_key=key)
            except Exception:
                return None
        return self._client

    def _set_profile(self, profile):
        """Update profile reference."""
        self.profile = profile

    def _on_config(self):
        """Re-create client after settings change."""
        self._client = None
        self._refresh_banner()

    def _refresh_banner(self):
        """Show red banner if API key missing."""
        if not self.cfg.get("api_key", "").strip():
            self.banner.setText(
                "⚠ No API key set — add one in Settings to enable "
                "CV generation.")
            self.banner.setVisible(True)
        else:
            self.banner.setVisible(False)

    def _apply_theme(self):
        """Apply dark theme stylesheet."""
        self.setStyleSheet(f"""
            QMainWindow, QWidget {{ background:{BG}; color:{FG}; }}
            QTabWidget::pane {{ border:1px solid {BORDER}; }}
            QTabBar::tab {{
                background:{PANEL}; color:{FG};
                padding:8px 20px; border:1px solid {BORDER};
            }}
            QTabBar::tab:selected {{ background:{BG}; border-bottom:none; }}
            QLineEdit, QTextEdit, QListWidget, QTableWidget, QComboBox {{
                background:{PANEL}; color:{FG}; border:1px solid {BORDER};
                padding:4px; selection-background-color:{ACCENT};
            }}
            QPushButton {{
                background:{PANEL}; color:{FG}; border:1px solid {BORDER};
                padding:6px 14px; border-radius:3px;
            }}
            QPushButton:hover {{ background:#21262d; }}
            QPushButton:disabled {{ color:#6e7681; }}
            QGroupBox {{
                border:1px solid {BORDER}; margin-top:8px; padding-top:8px;
            }}
            QGroupBox::title {{
                subcontrol-origin:margin; left:10px; padding:0 4px;
            }}
            QProgressBar {{
                background:{PANEL}; border:1px solid {BORDER};
                text-align:center; color:{FG};
            }}
            QProgressBar::chunk {{ background:{ACCENT}; }}
            QHeaderView::section {{
                background:{PANEL}; color:{FG}; padding:4px;
                border:1px solid {BORDER};
            }}
            QLabel {{ color:{FG}; }}
        """)


def main():
    """Entry point."""
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
