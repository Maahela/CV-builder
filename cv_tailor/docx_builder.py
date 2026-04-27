"""Render a tailored CV JSON to a formatted .docx file."""
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .constants import (ALWAYS_SHOW_SKILLS, CONDITIONAL_SKILLS, MARGIN_CM,
                        PAGE_H_CM, PAGE_W_CM, RIGHT_TAB_CM, SKILL_LABELS)


class DocxBuilder:
    """Produce a formatted .docx CV from tailored JSON data."""

    @staticmethod
    def _set_bottom_border(paragraph):
        """Add a single-line bottom border to a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _clear_table_borders(table):
        """Force all borders off (top/left/bottom/right/insideH/insideV)."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tblPr.append(borders)

    @staticmethod
    def _set_cell_width(cell, cm):
        """Set a fixed cell width in OXML (python-docx cell.width is lossy)."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(int(cm * 567)))  # 567 twips per cm
        tcW.set(qn("w:type"), "dxa")

    @staticmethod
    def _add_run(paragraph, text, *, bold=False, italic=False,
                 size=10.5, color=None, font="Calibri"):
        """Append a styled run to a paragraph."""
        run = paragraph.add_run(text)
        run.font.name = font
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rFonts.set(qn("w:cs"), font)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        return run

    @staticmethod
    def _set_defaults(doc):
        """Set document default font to Calibri 10.5 so styles inherit."""
        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        rPr = normal.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        rFonts.set(qn("w:cs"), "Calibri")

    @staticmethod
    def _set_page(doc):
        """Set A4, 2cm margins."""
        for section in doc.sections:
            section.page_width = Cm(PAGE_W_CM)
            section.page_height = Cm(PAGE_H_CM)
            section.left_margin = Cm(MARGIN_CM)
            section.right_margin = Cm(MARGIN_CM)
            section.top_margin = Cm(MARGIN_CM)
            section.bottom_margin = Cm(MARGIN_CM)

    @classmethod
    def _section_header(cls, doc, text):
        """Add an ALL CAPS section header with bottom border."""
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        cls._add_run(p, text.upper(), bold=True, size=11, color="000000")
        cls._set_bottom_border(p)
        return p

    @classmethod
    def _title_with_date(cls, doc, title, date_str, italic_title=False):
        """Bold title left, date right, right-tab at (page − margins)."""
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(RIGHT_TAB_CM), WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        cls._add_run(p, title, bold=True, italic=italic_title, size=11)
        if date_str:
            p.add_run("\t")
            cls._add_run(p, date_str, size=10)
        return p

    @classmethod
    def _subline(cls, doc, text, *, italic=True, color="444444"):
        """Italic subline (company / tech stack)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        cls._add_run(p, text, italic=italic, size=10.5, color=color)
        return p

    @classmethod
    def _bullets(cls, doc, items):
        """Add bullet list using Word's List Bullet style."""
        added = False
        for b in items:
            if not b:
                continue
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            cls._add_run(p, b, size=10.5)
            added = True
        if added:
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    @staticmethod
    def _add_hyperlink(paragraph, text, url, size=10):
        """Append a blue underlined clickable hyperlink run to a paragraph."""
        r_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/"
            "relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "1155CC")
        rPr.append(color_el)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        rPr.append(fonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))  # half-points
        rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    @classmethod
    def _header_block(cls, doc, profile):
        """Name + contact line (with clickable links) + hr."""
        name = profile.get("name", "") or ""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        cls._add_run(p, name, bold=True, size=18)

        contact = profile.get("contact", {}) or {}
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)

        def _plain(text):
            r = p2.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        def _sep():
            _plain(" | ")

        plain_fields = [
            contact.get("email", ""),
            contact.get("phone", ""),
        ]
        link_fields = {
            "linkedin": contact.get("linkedin", ""),
            "github":   contact.get("github", ""),
            "website":  contact.get("website", ""),
        }

        parts_plain = [f for f in plain_fields if f]
        for i, val in enumerate(parts_plain):
            _plain(val)
            if i < len(parts_plain) - 1 or any(link_fields.values()):
                _sep()

        link_items = [(v, v) for v in link_fields.values() if v]
        for i, (display, raw_url) in enumerate(link_items):
            url = raw_url if raw_url.startswith("http") else f"https://{raw_url}"
            cls._add_hyperlink(p2, display, url)
            if i < len(link_items) - 1:
                _sep()

        cls._set_bottom_border(p2)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(6)

    @classmethod
    def _skills_table(cls, doc, skills):
        """Two-column borderless table of skill categories."""
        ordered = ALWAYS_SHOW_SKILLS + CONDITIONAL_SKILLS
        rows = [(SKILL_LABELS[key], skills[key])
                for key in ordered
                if skills.get(key)]
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=2)
        table.autofit = False
        cls._clear_table_borders(table)
        tblPr = table._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)

        left_w = 3.5
        right_w = PAGE_W_CM - 2 * MARGIN_CM - left_w
        for row_i, (label, values) in enumerate(rows):
            left = table.cell(row_i, 0)
            right = table.cell(row_i, 1)
            cls._set_cell_width(left, left_w)
            cls._set_cell_width(right, right_w)
            lp = left.paragraphs[0]
            lp.paragraph_format.space_before = Pt(1)
            lp.paragraph_format.space_after = Pt(1)
            cls._add_run(lp, f"{label}:", bold=True, size=10.5)
            rp = right.paragraphs[0]
            rp.paragraph_format.space_before = Pt(1)
            rp.paragraph_format.space_after = Pt(1)
            cls._add_run(rp, ", ".join(values), size=10.5)

    @staticmethod
    def _remove_compat_mode(doc):
        """Strip the w:compat block so Word opens in edit mode, not compat mode."""
        settings = doc.settings.element
        compat = settings.find(qn("w:compat"))
        if compat is not None:
            settings.remove(compat)
        doc.core_properties.revision = 1

    @classmethod
    def build(cls, profile, cv_data, output_path):
        """Build the full CV document."""
        doc = Document()
        cls._remove_compat_mode(doc)
        cls._set_defaults(doc)
        cls._set_page(doc)

        name = cv_data.get("name") or profile.get("name", "")
        contact = cv_data.get("contact") or profile.get("contact", {})
        cls._header_block(doc, {"name": name, "contact": contact})

        summary = cv_data.get("summary") or profile.get("summary")
        if summary:
            cls._section_header(doc, "Summary")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            cls._add_run(p, summary, size=10.5)

        experience = cv_data.get("experience") or []
        if experience:
            cls._section_header(doc, "Experience")
            for role in experience:
                dates = cls._fmt_range(
                    role.get("start_date") or role.get("start"),
                    role.get("end_date") or role.get("end"),
                    role.get("current"))
                cls._title_with_date(doc, role.get("title", ""), dates)
                comp = role.get("company", "")
                loc = role.get("location", "")
                sub = comp + (f" — {loc}" if loc else "")
                if sub.strip():
                    cls._subline(doc, sub)
                bullets = (role.get("bullets")
                           or role.get("responsibilities")
                           or role.get("achievements") or [])
                if not bullets:
                    print(f"[warn] empty bullets for role: {role.get('title')}")
                cls._bullets(doc, bullets)

        def _norm(s):
            return re.sub(r"[^a-z0-9]+", " ", s.lower()).strip()

        projects = cv_data.get("projects") or []
        profile_proj_links = {
            _norm(p.get("name", "")): (p.get("link") or "").strip()
            for p in (profile.get("projects") or [])
        }
        if projects:
            cls._section_header(doc, "Projects")
            for proj in projects:
                date_str = cls._fmt_range(
                    proj.get("start_date"), proj.get("end_date"),
                    None) or str(proj.get("year") or "")
                proj_name = proj.get("name", "")
                proj_link = (proj.get("link") or "").strip() \
                    or profile_proj_links.get(_norm(proj_name), "")
                cls._title_with_date(doc, proj_name, date_str)
                tech = proj.get("technologies") or proj.get("tech") or []
                if isinstance(tech, list):
                    tech = ", ".join(tech)
                if tech:
                    cls._subline(doc, tech)
                if proj_link.startswith("https://"):
                    p_link = doc.add_paragraph()
                    p_link.paragraph_format.space_before = Pt(0)
                    p_link.paragraph_format.space_after = Pt(2)
                    cls._add_run(p_link, "GitHub: ", size=9, color="1155CC")
                    cls._add_hyperlink(p_link, proj_link, proj_link, size=9)
                bullets = proj.get("highlights") or proj.get("bullets") or []
                if not bullets:
                    desc = proj.get("description") or ""
                    if desc:
                        bullets = [s.strip() for s in re.split(r"\.\s+", desc)
                                   if s.strip()]
                if not bullets:
                    print(f"[warn] empty bullets for project: "
                          f"{proj.get('name')}")
                cls._bullets(doc, bullets)

        cv_skills = cv_data.get("skills") or {}
        profile_skills = profile.get("skills") or {}
        final_skills = {}
        for key in ALWAYS_SHOW_SKILLS:
            vals = cv_skills.get(key) or profile_skills.get(key) or []
            if vals:
                final_skills[key] = vals
        for key in CONDITIONAL_SKILLS:
            vals = cv_skills.get(key) or []
            if vals:
                final_skills[key] = vals
        if final_skills:
            cls._section_header(doc, "Technical Skills")
            cls._skills_table(doc, final_skills)

        volunteering = cv_data.get("volunteering") or []
        if volunteering:
            cls._section_header(doc, "Volunteering & Leadership")
            for v in volunteering:
                dates = cls._fmt_range(v.get("start_date"),
                                       v.get("end_date"), None) \
                        or v.get("period", "")
                cls._title_with_date(doc, v.get("role", ""), dates)
                org = v.get("organization") or v.get("org") or ""
                if org:
                    cls._subline(doc, org)
                bullets = v.get("bullets") or v.get("responsibilities") or []
                if not bullets:
                    desc = v.get("description") or ""
                    if desc:
                        bullets = [s.strip() for s in re.split(r"\.\s+", desc)
                                   if s.strip()]
                if not bullets:
                    print(f"[warn] empty bullets for volunteering: "
                          f"{v.get('role')}")
                cls._bullets(doc, bullets)

        achievements = cv_data.get("achievements") or []
        if achievements:
            cls._section_header(doc, "Achievements")
            cls._bullets(doc, achievements)

        education = cv_data.get("education") or profile.get("education") or []
        if education:
            cls._section_header(doc, "Education")
            for ed in education:
                dates = cls._fmt_range(ed.get("start_date"),
                                       ed.get("end_date"), None) \
                        or str(ed.get("year") or "")
                cls._title_with_date(doc, ed.get("degree", ""), dates)
                inst = ed.get("institution", "")
                if inst:
                    cls._subline(doc, inst)
                details = ed.get("details")
                if details:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(4)
                    cls._add_run(p, details, size=10, color="444444")

        doc.save(output_path)

    @staticmethod
    def _fmt_range(start, end, current):
        """Format a date range. Accepts start_date/end_date or start/end."""
        if current:
            end = "Present"
        start = (start or "").strip() if isinstance(start, str) else start
        end = (end or "").strip() if isinstance(end, str) else end
        if start and end:
            return f"{start} — {end}"
        return start or end or ""
